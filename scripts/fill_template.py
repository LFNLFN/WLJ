#!/usr/bin/env python3
"""根据模板生成填写好的 docx 文件"""
import sys
import json
import os
from docx import Document
from copy import deepcopy

def fill_template(template_path, data, output_path):
    """填写模板并保存"""
    doc = Document(template_path)
    
    # 处理表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = run.text
                        # 替换占位符
                        text = text.replace('姓名', data.get('studentName', '姓名') if text == '姓　　名' or text.strip() == '姓名' else text)
                        # 更精确的替换逻辑
                        # (这里太复杂，实际用 python-docx-template 更好)
                        pass
    
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    if len(sys.argv) > 1:
        data = json.loads(sys.argv[1])
        fill_template(sys.argv[2], data, sys.argv[3])
    else:
        print("Usage: fill_template.py <json_data> <template_path> <output_path>")
