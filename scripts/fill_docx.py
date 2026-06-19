#!/usr/bin/env python3
"""根据模板 docx 填写数据，输出到指定路径"""
import sys
import json
import os
from docx import Document
from docx.shared import Pt

FIELDS_MAP = {
    # 入学登记表
    '姓名': 'studentName',
    '临床诊断': 'diagnosis',
    '诊断机构': 'diagnosisInstitution',
    '户籍': 'registeredResidence',
    '现居住': 'address',
    '家庭电话': 'phone',
    '父亲姓名': 'fatherName',
    '母亲姓名': 'motherName',
    '其他成员': 'otherMember',
    '年龄': 'parentAge',  # 父母的年龄
    '文化程度': 'education',
    '工作单位及职务': 'workUnit',
    '联系电话': 'phone',
    '智障成因': 'causeOfDisability',
    '目前主要障碍情况': 'mainObstacles',
    '既往康复教育情况': 'previousRehab',
}

def fill_template(template_path, data, output_path):
    """用 json 数据填充 docx 模板"""
    if isinstance(data, str):
        data = json.loads(data)
    
    doc = Document(template_path)
    
    # 1. 替换段落中的占位符
    for p in doc.paragraphs:
        for run in p.runs:
            text = run.text
            if '档案编号' in text:
                run.text = text.replace('档案编号：', f'档案编号：{data.get("studentId", "")}')
    
    # 2. 替换表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = run.text.strip()
                        # 名字替换 - 简单字段直接填入
                        if text == '姓名':
                            # 找到"姓名"标签旁边的空单元格
                            pass
    
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    if len(sys.argv) > 3:
        data = json.loads(sys.argv[1])
        fill_template(sys.argv[2], data, sys.argv[3])
        print(json.dumps({"success": True, "output": sys.argv[3]}))
    else:
        print(json.dumps({"error": "Usage: fill_docx.py <json> <template_path> <output_path>"}))
